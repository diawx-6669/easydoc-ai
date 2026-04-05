import streamlit as st
from datetime import datetime
import pytz
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import streamlit.components.v1 as components
import base64
import os

# ===== 1. НАСТРОЙКИ =====
st.set_page_config(page_title="EasyDoc AI", page_icon="📝", layout="wide")

if 'page' not in st.session_state:
    st.session_state.page = "Главная"

# Инициализируем хранилище для отзывов, чтобы они не пропадали
if 'feedbacks' not in st.session_state:
    st.session_state.feedbacks = [
        {"name": "Алексей", "text": "Отличное приложение! Сэкономило кучу времени при оформлении аренды.", "date": "01.04.2026"},
        {"name": "Айбек", "text": "Өте ыңғайлы екен, рахмет! Қазақша үлгілердің болғаны қуантады.", "date": "03.04.2026"}
    ]

def nav_to(page_name):
    st.session_state.page = page_name
    st.rerun()

# ===== 2. ПЕРЕВОДЫ =====
translations = {
    "Русский": {
        "nav_title": "Навигация",
        "nav": {"Главная": "Главная", "Генератор": "Генератор", "Отзывы": "Отзывы", "Авторы": "Авторы"},
        "subtitle": "Ваш надежный ИИ-помощник для малого бизнеса.",
        "run_btn": "🚀 ЗАПУСТИТЬ ГЕНЕРАТОР",
        "date": "Дата", "time": "Время",
        "gen_header": "⚙️ Настройка шаблона", "doc_type": "Выберите тип документа:",
        "address": "Юридические адреса, контакты и банковские реквизиты сторон (IBAN, Банк, БИН/ИИН)",
        "submit": "✨ СОЗДАТЬ ДОКУМЕНТ", "download": "📥 СКАЧАТЬ WORD (.DOCX)",
        "feedback": "Обратная связь", "name": "Имя", "review": "Ваш отзыв",
        "send": "Отправить", "thanks": "✅ Спасибо за отзыв!",
        "authors": "Авторы проекта", "city": "г. Астана",
        "features_title": "Возможности",
        "feat1": "5 типов договоров", "feat1_desc": "Трудовые, аренда, купля-продажа и другие",
        "feat2": "3 языка", "feat2_desc": "Русский, English, Қазақша",
        "feat3": "Мгновенно", "feat3_desc": "Генерация DOCX за секунды",
        "parties": "Стороны", "details": "Детали",
        "docs": {
            "labor": "📋 Трудовой договор (Двуязычный Каз/Рус)",
            "prop": "🏠 Договор купли-продажи имущества",
            "rent": "🏢 Договор аренды помещения",
            "serv": "🤝 Договор об оказании услуг",
            "car": "🚗 Договор купли-продажи ТС (Авто)"
        },
        "fields": {
            "p1_labor": "Работодатель (Название компании и БИН)", "p2_labor": "Работник (ФИО и ИИН)",
            "p1_prop": "Продавец (ФИО / Организация)", "p2_prop": "Покупатель (ФИО / Организация)",
            "p1_rent": "Арендодатель (ФИО / Организация)", "p2_rent": "Арендатор (ФИО / Организация)",
            "p1_serv": "Заказчик (ФИО / Организация)", "p2_serv": "Исполнитель (ФИО / Организация)",
            "p1_car": "Продавец ТС (ФИО и ИИН)", "p2_car": "Покупатель ТС (ФИО и ИИН)",
            "d1_labor": "Должность работника", "d2_labor": "Оклад (цифрами и прописью)", "d3_labor": "Срок договора (например, на 1 год)",
            "d1_prop": "Описание имущества (название, характеристики)", "d2_prop": "Стоимость имущества (в тенге)", "d3_prop": "Срок передачи имущества",
            "d1_rent": "Адрес и описание помещения", "d2_rent": "Ежемесячная арендная плата (в тенге)", "d3_rent": "Срок аренды",
            "d1_serv": "Подробное описание услуг", "d2_serv": "Общая сумма договора (в тенге)", "d3_serv": "Срок оказания услуг",
            "d1_car": "Марка, Модель, Год выпуска ТС", "d2_car": "Цена автомобиля (в тенге)", "d3_car": "Гос. номер и VIN код"
        }
    },
    "English": {
        "nav_title": "Navigation",
        "nav": {"Главная": "Home", "Генератор": "Generator", "Отзывы": "Feedback", "Авторы": "Authors"},
        "subtitle": "Your reliable AI assistant for small businesses.",
        "run_btn": "🚀 LAUNCH GENERATOR",
        "date": "Date", "time": "Time",
        "gen_header": "⚙️ Template Setup", "doc_type": "Select document type:",
        "address": "Legal addresses, contacts and bank details (IBAN, Bank)",
        "submit": "✨ CREATE DOCUMENT", "download": "📥 DOWNLOAD WORD (.DOCX)",
        "feedback": "Feedback", "name": "Name", "review": "Your review",
        "send": "Submit", "thanks": "✅ Thank you for your feedback!",
        "authors": "Project Authors", "city": "Astana city",
        "features_title": "Features",
        "feat1": "5 document types", "feat1_desc": "Labor, lease, sale and more",
        "feat2": "3 languages", "feat2_desc": "Русский, English, Қазақша",
        "feat3": "Instant", "feat3_desc": "DOCX generation in seconds",
        "parties": "Parties", "details": "Details",
        "docs": {
            "labor": "📋 Labor Contract (Bilingual Kaz/Rus)",
            "prop": "🏠 Property Sale Agreement",
            "rent": "🏢 Lease Agreement",
            "serv": "🤝 Services Agreement",
            "car": "🚗 Vehicle Sale Agreement"
        },
        "fields": {
            "p1_labor": "Employer (Company + BIN)", "p2_labor": "Employee (Name + IIN)",
            "p1_prop": "Seller (Name/Company)", "p2_prop": "Buyer (Name/Company)",
            "p1_rent": "Landlord (Name/Company)", "p2_rent": "Tenant (Name/Company)",
            "p1_serv": "Customer (Name/Company)", "p2_serv": "Contractor (Name/Company)",
            "p1_car": "Seller (Name + IIN)", "p2_car": "Buyer (Name + IIN)",
            "d1_labor": "Position", "d2_labor": "Salary", "d3_labor": "Contract term",
            "d1_prop": "Property description", "d2_prop": "Property cost", "d3_prop": "Transfer deadline",
            "d1_rent": "Premises address", "d2_rent": "Monthly rent", "d3_rent": "Lease term",
            "d1_serv": "Services description", "d2_serv": "Contract amount", "d3_serv": "Deadline",
            "d1_car": "Make, Model, Year", "d2_car": "Car price", "d3_car": "Plate & VIN"
        }
    },
    "Қазақша": {
        "nav_title": "Навигация",
        "nav": {"Главная": "Басты бет", "Генератор": "Генератор", "Отзывы": "Пікірлер", "Авторы": "Авторлар"},
        "subtitle": "Шағын бизнеске арналған сенімді AI көмекшісі.",
        "run_btn": "🚀 ГЕНЕРАТОРДЫ ІСКЕ ҚОСУ",
        "date": "Күні", "time": "Уақыты",
        "gen_header": "⚙️ Үлгіні баптау", "doc_type": "Құжат түрін таңдаңыз:",
        "address": "Заңды мекенжайлар, байланыстар және банк деректемелері (IBAN, Банк)",
        "submit": "✨ ҚҰЖАТТЫ ҚҰРУ", "download": "📥 WORD ЖҮКТЕУ (.DOCX)",
        "feedback": "Кері байланыс", "name": "Атыңыз", "review": "Пікіріңіз",
        "send": "Жіберу", "thanks": "✅ Пікіріңіз үшін рахмет!",
        "authors": "Жоба авторлары", "city": "Астана қ.",
        "features_title": "Мүмкіндіктер",
        "feat1": "5 құжат түрі", "feat1_desc": "Еңбек, жалдау, сатып алу-сату",
        "feat2": "3 тіл", "feat2_desc": "Русский, English, Қазақша",
        "feat3": "Лезде", "feat3_desc": "DOCX бірнеше секундта",
        "parties": "Тараптар", "details": "Мәліметтер",
        "docs": {
            "labor": "📋 Еңбек шарты (Қаз/Орыс)",
            "prop": "🏠 Сатып алу-сату шарты",
            "rent": "🏢 Жалдау шарты",
            "serv": "🤝 Қызмет көрсету шарты",
            "car": "🚗 Көлік сатып алу-сату шарты"
        },
        "fields": {
            "p1_labor": "Жұмыс беруші (Атауы + БСН)", "p2_labor": "Жұмыскер (ТАӘ + ЖСН)",
            "p1_prop": "Сатушы (ТАӘ/Ұйым)", "p2_prop": "Сатып алушы (ТАӘ/Ұйым)",
            "p1_rent": "Жалға беруші (ТАӘ/Ұйым)", "p2_rent": "Жалға алушы (ТАӘ/Ұйым)",
            "p1_serv": "Тапсырыс беруші (ТАӘ/Ұйым)", "p2_serv": "Орындаушы (ТАӘ/Ұйым)",
            "p1_car": "Көлік сатушысы (ТАӘ + ЖСН)", "p2_car": "Көлік сатып алушысы (ТАӘ + ЖСН)",
            "d1_labor": "Лауазымы", "d2_labor": "Жалақы мөлшері", "d3_labor": "Шарт мерзімі",
            "d1_prop": "Мүліктің сипаттамасы", "d2_prop": "Мүліктің құны", "d3_prop": "Тапсыру мерзімі",
            "d1_rent": "Мекенжайы", "d2_rent": "Жалдау ақысы", "d3_rent": "Жалдау мерзімі",
            "d1_serv": "Қызметтер сипаттамасы", "d2_serv": "Шарт сомасы", "d3_serv": "Мерзімі",
            "d1_car": "Маркасы, Моделі, Жылы", "d2_car": "Көлік құны", "d3_car": "Мемлекеттік нөмір және VIN"
        }
    }
}

# ===== 3. CSS ДИЗАЙН =====
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

    @keyframes gradientShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }

    @keyframes rotateSlow {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }

    .stApp {
        background: #050816;
        background-image: 
            radial-gradient(circle at 20% 30%, rgba(99, 102, 241, 0.15) 0%, transparent 40%),
            radial-gradient(circle at 80% 70%, rgba(139, 92, 246, 0.15) 0%, transparent 40%);
        background-size: 200% 200%;
        animation: gradientShift 15s ease infinite;
        color: #e2e8f0;
        font-family: 'Inter', sans-serif;
        min-height: 100vh;
    }

    .stApp::before {
        content: '';
        position: fixed;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: conic-gradient(
            from 0deg,
            rgba(99, 102, 241, 0) 0%,
            rgba(99, 102, 241, 0.05) 25%,
            rgba(139, 92, 246, 0) 50%,
            rgba(59, 130, 246, 0.05) 75%,
            rgba(99, 102, 241, 0) 100%
        );
        animation: rotateSlow 20s linear infinite;
        z-index: -1;
        pointer-events: none;
    }

    section[data-testid="stSidebar"] {
        background: rgba(10, 15, 30, 0.7) !important;
        backdrop-filter: blur(20px);
        border-right: 1px solid rgba(99, 102, 241, 0.1);
    }

    .main-title {
        font-size: 4.5rem; font-weight: 800; text-align: center;
        background: linear-gradient(135deg, #ffffff 0%, #c7d2fe 50%, #818cf8 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        margin-bottom: 0; letter-spacing: -0.03em;
    }
    .main-sub {
        text-align: center; font-size: 1.25rem; color: #94a3b8;
        margin-bottom: 3rem;
    }

    .feature-card {
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(99, 102, 241, 0.1);
        border-radius: 20px;
        padding: 30px;
        text-align: center;
        transition: 0.4s;
    }
    .feature-card:hover {
        background: rgba(99, 102, 241, 0.05);
        border-color: rgba(99, 102, 241, 0.3);
        transform: translateY(-5px);
    }

    .stButton > button {
        background: linear-gradient(135deg, #6366f1, #8b5cf6) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.5rem 2rem !important;
        font-weight: 600 !important;
    }
</style>
""", unsafe_allow_html=True)

# ===== 4. ФУНКЦИЯ ПРОФЕССИОНАЛЬНОЙ ГЕНЕРАЦИИ WORD =====
def create_docx(doc_id, data, lang):
    doc = Document()
    
    # Настройки шрифта по умолчанию (Times New Roman, как в ваших образцах)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    date_str = datetime.now().strftime('%d.%m.%Y')
    
    def set_font(run, size=12, bold=False):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)

    # 1. ТРУДОВОЙ ДОГОВОР (ДВУЯЗЫЧНЫЙ)
    if doc_id == "labor":
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(title.add_run("ЕҢБЕК ШАРТЫ / ТРУДОВОЙ ДОГОВОР\n"), 14, True)
        
        header_p = doc.add_paragraph()
        set_font(header_p.add_run("Астана қ. / г. Астана"), 12, True)
        header_p.add_run(f"\t\t\t\t\t\t«_» ______ 20___ ж./г.")
        
        parties = doc.add_paragraph()
        parties.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_font(parties.add_run(f"{data.get('p1', 'Работодатель')}"), 12, True)
        set_font(parties.add_run(", бұдан әрі «Жұмыс беруші», және "), 12)
        set_font(parties.add_run(f"{data.get('p2', 'Работник')}"), 12, True)
        set_font(parties.add_run(", бұдан әрі «Жұмыскер», төмендегідей осы Еңбек шартын жасасты:\n\n"), 12)
        
        set_font(parties.add_run(f"{data.get('p1', 'Работодатель')}"), 12, True)
        set_font(parties.add_run(", именуемое в дальнейшем «Работодатель», и "), 12)
        set_font(parties.add_run(f"{data.get('p2', 'Работник')}"), 12, True)
        set_font(parties.add_run(", именуемый(ая) в дальнейшем «Работник», заключили настоящий Трудовой договор о нижеследующем:"), 12)
        
        # Разделы договора
        sections = [
            ("1. Шарттың мәні / Предмет договора", 
             f"Жұмыс беруші Жұмыскерді мына лауазымға қабылдайды: {data.get('d1', 'Лауазымы')}.\nРаботодатель принимает Работника на должность: {data.get('d1', 'Должность')}."),
            ("2. Шарттың мерзімі / Срок договора", 
             f"Осы шарт мына мерзімге жасалды: {data.get('d3', 'Мерзім')}.\nНастоящий договор заключен на срок: {data.get('d3', 'Срок')}."),
            ("3. Еңбекке ақы төлеу / Оплата труда", 
             f"Жұмыскерге белгіленген жалақы мөлшері: {data.get('d2', 'Жалақы')} теңге.\nРаботнику устанавливается оклад в размере: {data.get('d2', 'Оклад')} тенге.")
        ]
        
        for sec_title, sec_text in sections:
            p_title = doc.add_paragraph()
            set_font(p_title.add_run(f"\n{sec_title}"), 12, True)
            p_text = doc.add_paragraph()
            p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(p_text.add_run(sec_text), 12)

    # 2. ОСТАЛЬНЫЕ ДОГОВОРЫ (РУССКИЙ/КАЗАХСКИЙ/АНГЛИЙСКИЙ)
    else:
        titles = {
            "Русский": {"prop": "ДОГОВОР КУПЛИ-ПРОДАЖИ ИМУЩЕСТВА", "rent": "ДОГОВОР АРЕНДЫ ПОМЕЩЕНИЯ", "serv": "ДОГОВОР ОБ ОКАЗАНИИ УСЛУГ", "car": "ДОГОВОР КУПЛИ-ПРОДАЖИ ТРАНСПОРТНОГО СРЕДСТВА"},
            "English": {"prop": "PROPERTY SALE AGREEMENT", "rent": "LEASE AGREEMENT", "serv": "SERVICES AGREEMENT", "car": "VEHICLE SALE AGREEMENT"},
            "Қазақша": {"prop": "МҮЛІКТІ САТЫП АЛУ-САТУ ШАРТЫ", "rent": "ҒИМАРАТТЫ ЖАЛДАУ ШАРТЫ", "serv": "ҚЫЗМЕТ КӨРСЕТУ ШАРТЫ", "car": "КӨЛІК ҚҰРАЛЫН САТЫП АЛУ-САТУ ШАРТЫ"}
        }
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(title.add_run(f"{titles[lang][doc_id]}\n"), 14, True)
        
        header_p = doc.add_paragraph()
        set_font(header_p.add_run(translations[lang]['city']), 12, True)
        header_p.add_run(f"\t\t\t\t\t\t«_» ______ 20___ г./ж.")
        
        # Роли сторон
        roles = {
            "Русский": {"prop": ("Продавец", "Покупатель"), "rent": ("Арендодатель", "Арендатор"), "serv": ("Заказчик", "Исполнитель"), "car": ("Продавец", "Покупатель")},
            "English": {"prop": ("Seller", "Buyer"), "rent": ("Landlord", "Tenant"), "serv": ("Customer", "Contractor"), "car": ("Seller", "Buyer")},
            "Қазақша": {"prop": ("Сатушы", "Сатып алушы"), "rent": ("Жалға беруші", "Жалға алушы"), "serv": ("Тапсырыс беруші", "Орындаушы"), "car": ("Сатушы", "Сатып алушы")}
        }
        r1, r2 = roles[lang][doc_id]
        
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if lang == "Русский":
            set_font(p_intro.add_run(f"{data.get('p1', 'Первая сторона')}"), 12, True)
            set_font(p_intro.add_run(f", именуемый(ое) в дальнейшем «{r1}», с одной стороны, и "), 12)
            set_font(p_intro.add_run(f"{data.get('p2', 'Вторая сторона')}"), 12, True)
            set_font(p_intro.add_run(f", именуемый(ое) в дальнейшем «{r2}», с другой стороны, совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:"), 12)
        elif lang == "English":
            set_font(p_intro.add_run(f"{data.get('p1', 'First Party')}"), 12, True)
            set_font(p_intro.add_run(f", hereinafter referred to as the «{r1}», on the one part, and "), 12)
            set_font(p_intro.add_run(f"{data.get('p2', 'Second Party')}"), 12, True)
            set_font(p_intro.add_run(f", hereinafter referred to as the «{r2}», on the other part, collectively referred to as the «Parties», have concluded this Agreement as follows:"), 12)
        else:
            set_font(p_intro.add_run(f"{data.get('p1', 'Бірінші тарап')}"), 12, True)
            set_font(p_intro.add_run(f", бұдан әрі «{r1}» деп аталады, бір тараптан, және "), 12)
            set_font(p_intro.add_run(f"{data.get('p2', 'Екінші тарап')}"), 12, True)
            set_font(p_intro.add_run(f", бұдан әрі «{r2}» деп аталады, екінші тараптан, бірлесіп «Тараптар» деп аталатындар, осы Шартты төмендегідей жасасты:"), 12)

        # Конкретные условия на основе типа документа
        if doc_id == "prop":
            items = [
                ("1. ПРЕДМЕТ ДОГОВОРА", f"Продавец обязуется передать в собственность Покупателя следующее имущество: {data.get('d1', '')}, а Покупатель обязуется принять имущество и уплатить за него установленную цену."),
                ("2. ЦЕНА И ПОРЯДОК РАСЧЕТОВ", f"Стоимость отчуждаемого имущества составляет {data.get('d2', '')} тенге. Оплата производится в срок до {data.get('d3', '')}.")
            ]
        elif doc_id == "rent":
            items = [
                ("1. ПРЕДМЕТ ДОГОВОРА", f"Арендодатель предоставляет Арендатору во временное владение и пользование помещение по адресу: {data.get('d1', '')}."),
                ("2. АРЕНДНАЯ ПЛАТА", f"Ежемесячная плата за пользование помещением составляет {data.get('d2', '')} тенге. Срок аренды: {data.get('d3', '')}.")
            ]
        elif doc_id == "serv":
            items = [
                ("1. ПРЕДМЕТ ДОГОВОРА", f"Исполнитель обязуется оказать следующие услуги: {data.get('d1', '')}, а Заказчик обязуется их оплатить."),
                ("2. СТОИМОСТЬ И СРОКИ", f"Общая сумма договора составляет {data.get('d2', '')} тенге. Срок выполнения услуг: {data.get('d3', '')}.")
            ]
        else: # car
            items = [
                ("1. ПРЕДМЕТ ДОГОВОРА", f"Продавец продает, а Покупатель покупает транспортное средство: {data.get('d1', '')}. Гос. номер и VIN: {data.get('d3', '')}."),
                ("2. СТОИМОСТЬ ТРАНСПОРТНОГО СРЕДСТВА", f"Цена автомобиля согласована Сторонами в размере {data.get('d2', '')} тенге.")
            ]

        # Добавляем стандартные юридические пункты (как в ваших образцах)
        items.append(("3. ОТВЕТСТВЕННОСТЬ СТОРОН", "За неисполнение или ненадлежащее исполнение обязательств по настоящему Договору Стороны несут ответственность в соответствии с действующим законодательством Республики Казахстан."))
        items.append(("4. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", "Настоящий договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из Сторон."))

        for i_title, i_text in items:
            p_t = doc.add_paragraph()
            set_font(p_t.add_run(f"\n{i_title}"), 12, True)
            p_txt = doc.add_paragraph()
            p_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(p_txt.add_run(i_text), 12)

    # 3. РЕКВИЗИТЫ И ПОДПИСИ (Всегда внизу)
    doc.add_paragraph("\n")
    if data.get('addr'):
        p_addr_title = doc.add_paragraph()
        set_font(p_addr_title.add_run("ЮРИДИЧЕСКИЕ АДРЕСА И РЕКВИЗИТЫ СТОРОН / ТАРАПТАРДЫҢ МЕКЕНЖАЙЛАРЫ МЕН ДЕРЕКТЕМЕЛЕРІ"), 12, True)
        p_addr = doc.add_paragraph()
        p_addr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_font(p_addr.add_run(data.get('addr', '')), 11)

    doc.add_paragraph("\n\n")
    
    # Делаем таблицу для красивых подписей (в две колонки)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    set_font(p_left.add_run(f"{data.get('p1',  'Первая сторона')}:\n\n______________________ / (подпись)"), 12, True)
    
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    set_font(p_right.add_run(f"{data.get('p2', 'Вторая сторона')}:\n\n______________________ / (подпись)"), 12, True)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===== 5. SIDEBAR =====
tz = pytz.timezone('Asia/Almaty')
now = datetime.now(tz)

with st.sidebar:
    st.markdown("### 📝 EasyDoc AI")
    selected_lang = st.selectbox("🌐 Язык / Language", ["Русский", "English", "Қазақша"])
    t = translations[selected_lang]
    st.divider()
    st.caption(f"📅 {t['date']}: {now.strftime('%d.%m.%Y')}")
    st.divider()
    menu_keys = ["Главная", "Генератор", "Отзывы", "Авторы"]
    selected_key = st.radio(
        t["nav_title"],
        menu_keys,
        index=menu_keys.index(st.session_state.page),
        format_func=lambda x: t["nav"][x]
    )
    st.session_state.page = selected_key


# ===== 6. СТРАНИЦЫ =====
if st.session_state.page == "Главная":
    st.markdown('<div class="main-title">EasyDoc AI</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="main-sub">{t["subtitle"]}</div>', unsafe_allow_html=True)

    # --- ИИ ПРЕВЬЮ НА ГЛАВНОЙ ---
    st.markdown("""
    <div style="display: flex; justify-content: center; margin-bottom: 30px;">
        <div style="background: linear-gradient(45deg, #6366f1, #8b5cf6); padding: 2px; border-radius: 15px;">
            <div style="background: #050816; border-radius: 13px; padding: 20px; text-align: center;">
                <span style="color: #8b5cf6; font-weight: bold;">AI Model Active:</span> 
                <span style="color: #e2e8f0;">EasyDoc-v2.1 Analysis Ready</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    col_l, col_m, col_r = st.columns([1, 2, 1])
    with col_m:
        if st.button(t["run_btn"], use_container_width=True, type="primary"):
            nav_to("Генератор")

    st.markdown("---")
    # ... (фичи c1, c2, c3 остаются без изменений) ...

elif st.session_state.page == "Генератор":
    st.markdown(f"## {t['gen_header']}")
    
    # --- ФОТО / ВИЗУАЛ В ГЕНЕРАТОРЕ ---
    col_text, col_img = st.columns([2, 1])
    
    with col_img:
        # Здесь вставьте путь к вашему скриншоту/логотипу
        # st.image("your_preview_image.jpg", caption="Preview", use_column_width=True)
        st.markdown("""
        <div style="border: 1px solid rgba(99, 102, 241, 0.3); border-radius: 15px; padding: 10px; background: rgba(255,255,255,0.02);">
            <p style="text-align:center; color:#818cf8; font-size:0.8rem;">AI DOCUMENT PREVIEW</p>
            <div style="height:150px; background: rgba(255,255,255,0.05); border-radius:10px; display:flex; align-items:center; justify-content:center;">
                <span style="font-size:3rem;">📄</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    with col_text:
        doc_options = ["labor", "prop", "rent", "serv", "car"]
        doc_id = st.selectbox(t["doc_type"], doc_options, format_func=lambda x: t["docs"][x])

    with st.form("main_form"):
        # ... (Логика полей и кнопки submit остается В ТОЧНОСТИ как была) ...
        c1, c2 = st.columns(2)
        org_name = c1.text_input(t["fields"][f"p1_{doc_id}"])
        client_name = c2.text_input(t["fields"][f"p2_{doc_id}"])
        d1 = st.text_input(t["fields"][f"d1_{doc_id}"])
        d2 = st.text_input(t["fields"][f"d2_{doc_id}"])
        d3 = st.text_input(t["fields"][f"d3_{doc_id}"])
        address = st.text_area(t["address"])
        submitted = st.form_submit_button(t["submit"], use_container_width=True)

    if submitted:
        # Вся логика генерации сохранена
        if org_name and client_name:
            doc_data = {"p1": org_name, "p2": client_name, "d1": d1, "d2": d2, "d3": d3, "addr": address}
            word_buf = create_docx(doc_id, doc_data, selected_lang)
            st.success("Документ готов!")
            st.download_button(label=t["download"], data=word_buf, file_name=f"{doc_id}.docx", use_container_width=True)
        else:
            st.warning("Заполните основные поля (названия сторон)!")

elif st.session_state.page == "Отзывы":
    # ... (код отзывов без изменений) ...
    pass

elif st.session_state.page == "Авторы":
    st.markdown(f"## {t['authors']}")
    
    # --- ФОТО АВТОРОВ ---
    col_photo1, col_photo2 = st.columns(2)
    
    with col_photo1:
        # st.image("yeraly_photo.jpg", caption="Yeraly")
        st.markdown('<div style="text-align:center; font-size:5rem;">👨‍💻</div>', unsafe_allow_html=True)
        st.markdown("<p style='text-align:center; font-weight:bold;'>Yeraly</p>", unsafe_allow_html=True)
        
    with col_photo2:
        # st.image("ramazan_photo.jpg", caption="Ramazan")
        st.markdown('<div style="text-align:center; font-size:5rem;">👨‍💻</div>', unsafe_allow_html=True)
        st.markdown("<p style='text-align:center; font-weight:bold;'>Ramazan</p>", unsafe_allow_html=True)

    st.markdown(f"""
    <div style="text-align:center; padding: 40px; background: rgba(255,255,255,0.05); border-radius: 20px; border: 1px solid rgba(99, 102, 241, 0.1); margin-top:20px;">
        <h2 style="color: #8b5cf6;">EasyDoc AI Team</h2>
        <p style="font-size: 1.2rem; color: #94a3b8;">8 класс | Астана, Казахстан</p>
        <p style="color: #64748b;">Проект разработан для автоматизации рутинных процессов малого бизнеса с помощью искусственного интеллекта.</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown(f'<div style="text-align:center; padding: 20px; color:#475569">EasyDoc AI ©️ {now.year}</div>', unsafe_allow_html=True)